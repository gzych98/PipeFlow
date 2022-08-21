from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from matplotlib.pyplot import Figure
from tkinter import *
from subprocess import check_output
from scipy.signal import argrelextrema
import numpy as np
from docx import Document

path = 'D:\z03v22.exe'


def create_doc(dane):
    print(dane.params)
    spr = Document()
    spr.add_heading('Temat projektu', level=1)
    spr.add_heading('1. Badanie wstępne', level=1)
    spr.add_heading(f'Wielkość próby', level=2)
    spr.add_paragraph(f'\t{dane.params["wlk_prb"]} dla punktu {dane.params["pnkt"]}')
    spr.add_heading('Średnia', level=2)
    spr.add_paragraph(f'\t{dane.params["srednia"]}')
    spr.add_heading('Mediana', level=2)
    spr.add_paragraph(f'\t{dane.params["mediana"]}')
    spr.add_heading('Wariancja', level=2)
    spr.add_paragraph(f'\t{dane.params["wariacja"]}')
    spr.add_heading('Odchylenie standardowe', level=2)
    spr.add_paragraph(f'\t{dane.params["odch"]}')
    spr.add_heading('Minimum', level=2)
    spr.add_paragraph(f'\tbrak danych')
    spr.add_heading('Maksimum', level=2)
    spr.add_paragraph(f'\tbrak danych')
    spr.add_picture('plotMOiPE.png')
    spr.save('sprawko.docx')


def callmyexe(value):
    return float(check_output([path, str(value)]))


class StaticParameters:
    def __init__(self, range1, point1):
        self.range1 = range1
        self.point1 = point1
        self.params = {}
        self.a = np.full(self.range1, self.point1/1000)
        self.b = []
        for i in self.a:
            self.b.append(callmyexe(i))
        T2.insert(END, f'Wielkość próby: {self.range1}\ndla punktu pomiarowego: {self.point1}\n')
        self.params['wlk_prb'] = self.range1
        self.params['pnkt'] = self.point1

    def average(self, print_value):
        c = 0
        for i in self.b:
            c += i
        d = c / self.range1
        if print_value:
            T2.insert(END, f'Średnia: \n  {d}\n')
        else:
            return d
        self.params['srednia'] = d

    def mediana(self):
        self.b.sort()
        mediana1 = 0.0
        length = len(self.b)
        len_mode = length % 2
        if len_mode == 0:
            pos = int(length/2)
            mediana1 += (self.b[pos] + self.b[pos+1]) / 2
            T3.insert(END, F'Mediana:\n  {mediana1}\n')
        else:
            mediana1 += self.b[int((length + 1) / 2)]
            T3.insert(END, F'Mediana:\n  {mediana1}\n')
        self.params['mediana'] = mediana1

    def wariacja(self, print_value):
        avarage = StaticParameters.average(self, False)
        x = 0
        for i in self.b:
            x += ((i-avarage)**2)
        wariacja = x/len(self.b)
        if print_value:
            T3.insert(END, f'Wariancja:\n  {wariacja}\n')
        else:
            return wariacja
        self.params['wariacja'] = wariacja

    def odch_std(self):
        wariancja = StaticParameters.wariacja(self, False)
        odch_std = np.sqrt(wariancja)
        T4.insert(END, f'Odchylenie standardowe:\n  {odch_std}\n')
        self.params['odch'] = odch_std


def get_variable():
    try:
        try_static2 = StaticParameters(int(entry_variable2.get()), int(entry_variable3.get()))
        T1_extrema.delete(1.0, END)
        T2.delete(1.0, END)
        T3.delete(1.0, END)
        T4.delete(1.0, END)
        c = []
        for i in range(0, int(entry_variable1.get()) + 1, 1):
            a = i / 1000
            b = callmyexe(a)
            c.append(b)
        Chart.plot(chart1, c)
        Chart.plot2(chart2, try_static2.b)
        y = np.asarray(c)
        x = argrelextrema(y, np.greater)
        T1_extrema.insert(END, f'Ekstrema lokalne:\n')
        for i in x:
            for j in i:
                T1_extrema.insert(END, f'dla x = {j / 1000}\n  V = {c[j]}\n')
        del x
    except ValueError:
        print('That is NOT number')


class Chart:
    def __init__(self, frame1):
        self.frame1 = frame1

    def plot(self, a):
        try:
            self.object.get_tk_widget().pack_forget()
            self.object1.get_tk_widget().pack_forget()
        except AttributeError:
            pass
        f = Figure(figsize=(6, 4))
        g = Figure(figsize=(6, 4))
        d = np.arange(0.001, len(a) / 1000 + 0.001, 0.001)
        plot1 = f.add_subplot(111)
        plot2 = g.add_subplot(111)
        plot1.set_xlabel('Odległość od osi [m]')
        plot2.set_xlabel('Odległość od osi [m]')
        plot1.set_ylabel('Prędkość [m/s]')
        plot1.plot(d, a)
        plot2.set_yscale('log')
        plot2.plot(d, np.log10(a))
        self.object = FigureCanvasTkAgg(f, master=self.frame1)
        self.object.get_tk_widget().pack()
        self.object.draw()
        self.object1 = FigureCanvasTkAgg(g, master=self.frame1)
        self.object1.get_tk_widget().pack()
        self.object1.draw()
        try_static1 = StaticParameters(int(entry_variable2.get()), int(entry_variable3.get()))
        try_static1.average(True)
        try_static1.mediana()
        try_static1.wariacja(True)
        try_static1.odch_std()
        f.savefig('plotMOiPe.png', transparent=True)
        create_doc(try_static1)

    def plot2(self, a):
        try:
            self.object.get_tk_widget().pack_forget()
        except AttributeError:
            pass
        f = Figure(figsize=(6, 4))
        plot1 = f.add_subplot(111)
        y_pos = np.arange(1, len(a)+1)
        plot1.set_xlim([0, len(a)+1])
        plot1.bar(y_pos, a, color='#969696')
        self.object = FigureCanvasTkAgg(f, master=self.frame1)
        self.object.get_tk_widget().pack()
        self.object.draw()


root = Tk()
root.configure(bg='white')
root.geometry("1200x1000+0+0")
width_root = 1200
root.title("MOiPE Project 1")
# Frames
top_frame = Frame(root,
                  bg="#222",
                  width=width_root,
                  height=50)
lower_top_frame = Frame(root,
                        bg="#333",
                        width=width_root,
                        height=50)
show_info = Frame(root,
                  bg='#555',
                  width=width_root,
                  height=100)
top_frame.place(x=0,
                y=0)
lower_top_frame.place(x=0,
                      y=50)
aframe = Frame(root)
bfarme = Frame(root)
show_info.place(x=0,
                y=100)
# INSIDE FRAMES
T1_extrema = Text(show_info,
                  height=5,
                  width=18,
                  bg="#fff")
T1_extrema.place(x=10,
                 y=10,
                 anchor='nw')
T2 = Text(show_info,
          height=5,
          width=30)
T2.place(x=320,
         y=10,
         anchor='nw')
T3 = Text(show_info,
          height=5,
          width=30)
T3.place(x=570,
         y=10,
         anchor='nw')
T4 = Text(show_info,
          height=5,
          width=30)
T4.place(x=820,
         y=10,
         anchor='nw')

label_name = Label(top_frame,
                   text='Wyznaczenie profilu prędkości przepływu płynu w rurze o przekroju kołowym',
                   bg="#222",
                   font=("Verdana", 15),
                   fg='white')
label_name.place(x=width_root / 2,
                 y=25,
                 anchor="center")
label_try = Label(show_info,
                  text='Parametry statyczne \ndla wybranego punktu:',
                  fg='white',
                  bg='#555')
label_try.place(x=175,
                y=30)
label_variable = Label(lower_top_frame,
                       text='Wprowadź zmienne:',
                       font=("Verdana", 12),
                       fg='white',
                       bg="#333")
label_variable.place(x=20,
                     y=25,
                     anchor="w")
first_label = Label(lower_top_frame,
                    text='pierwsza:',
                    fg='white',
                    bg="#333",
                    )

last_label = Label(lower_top_frame,
                   text='promień wewnętrzny rury [mm]:',
                   fg='white',
                   bg="#333")
last_label.place(x=200,
                 y=25,
                 anchor="w")
range_label = Label(lower_top_frame,
                    text='wielkość próby:',
                    fg='white',
                    bg="#333")
range_label.place(x=460,
                  y=25,
                  anchor="w")
point_label = Label(lower_top_frame,
                    text='dla punktu:',
                    fg='white',
                    bg="#333")
point_label.place(x=630,
                  y=25,
                  anchor="w")
entry_variable = Entry(lower_top_frame, width=10)

entry_variable1 = Entry(lower_top_frame, width=10)
entry_variable1.place(x=385,
                      y=25,
                      anchor="w")
entry_variable2 = Entry(lower_top_frame, width=10)
entry_variable2.place(x=555,
                      y=25,
                      anchor="w")
entry_variable3 = Entry(lower_top_frame, width=10)
entry_variable3.place(x=700,
                      y=25,
                      anchor="w")

chart1 = Chart(aframe)
chart2 = Chart(bfarme)

button_variable = Button(lower_top_frame,
                         text='OK',
                         command=get_variable)
button_variable.place(x=1000,
                      y=22,
                      anchor="w")

aframe.place(x=0,
             y=200)
bfarme.place(x=600,
             y=200)

root.mainloop()
